"""DOCX 渲染。使用 python-docx。"""
from __future__ import annotations
from pathlib import Path
from typing import List, Any

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re


SEVERITY_LABEL = {'low': '低', 'medium': '中', 'high': '高'}


def _set_cn_font(run, size=10.5, bold=False):
    run.font.name = '微软雅黑'
    from docx.oxml.ns import qn
    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    if bold:
        run.bold = True


def _add_md_paragraphs(doc, text: str):
    """把 markdown 简单语法转成 Word 段落/列表/加粗。"""
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # 整行被 ** 包裹 → 当作加粗小标题
        m_heading = re.match(r'\*\*(.+?)\*\*$', line)
        if m_heading:
            p = doc.add_paragraph()
            _set_cn_font(p.add_run(m_heading.group(1)), size=12, bold=True)
            i += 1
            continue

        # 编号列表 1. / 2. 等
        m_num = re.match(r'^(\d+)\.\s+(.+)$', line)
        if m_num:
            p = doc.add_paragraph(style='List Number')
            _add_md_runs(p, m_num.group(2))
            i += 1
            continue

        # 项目符号 - / * / +
        m_bullet = re.match(r'^[-\*+]\s+(.+)$', line)
        if m_bullet:
            p = doc.add_paragraph(style='List Bullet')
            _add_md_runs(p, m_bullet.group(1))
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_md_runs(p, line)
        i += 1


def _add_md_runs(paragraph, text: str):
    """解析行内 **bold** 和 *italic*，分段添加 run。"""
    # 先处理 **bold**
    parts = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            _set_cn_font(paragraph.add_run(part[2:-2]), bold=True)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            _set_cn_font(run)
            run.italic = True
        else:
            _set_cn_font(paragraph.add_run(part))


def render_docx(report, detections: List[Any], output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(report.title)
    r.bold = True
    _set_cn_font(r, size=20)

    # 1. 基本信息
    doc.add_heading('一、基本信息', level=2)
    if report.lab_name:
        cells = [
            ('实验室', report.lab_name, '检查人', report.inspector or '-'),
            ('总体风险', SEVERITY_LABEL.get(report.overall_severity, '中'),
             '生成时间', report.created_at.strftime('%Y-%m-%d %H:%M')),
            ('创建人', report.created_by.username if report.created_by_id else '-',
             '隐患图片数', str(len(detections))),
        ]
        info_table = doc.add_table(rows=3, cols=4)
    elif report.address:
        cells = [
            ('地址', report.address, '检查人', report.inspector or '-'),
            ('总体风险', SEVERITY_LABEL.get(report.overall_severity, '中'),
             '生成时间', report.created_at.strftime('%Y-%m-%d %H:%M')),
            ('创建人', report.created_by.username if report.created_by_id else '-',
             '隐患图片数', str(len(detections))),
        ]
        info_table = doc.add_table(rows=3, cols=4)
    else:
        cells = [
            ('检查人', report.inspector or '-', '生成时间',
             report.created_at.strftime('%Y-%m-%d %H:%M')),
            ('创建人', report.created_by.username if report.created_by_id else '-',
             '隐患图片数', str(len(detections))),
        ]
        info_table = doc.add_table(rows=2, cols=4)
    info_table.style = 'Light Grid Accent 1'
    for ri, row in enumerate(cells):
        for ci, val in enumerate(row):
            cell = info_table.rows[ri].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            _set_cn_font(run)

    # 2. 汇总
    doc.add_heading('二、汇总统计', level=2)
    stats = report.summary_stats or {}
    by_sev = stats.get('by_severity', {}) or {}
    by_cat = stats.get('by_category', {}) or {}

    sev_p = doc.add_paragraph()
    _set_cn_font(sev_p.add_run(
        f"高:{by_sev.get('high',0)}  中:{by_sev.get('medium',0)}  低:{by_sev.get('low',0)}  "
        f"合计:{stats.get('total', 0)}"))

    if by_cat:
        cat_table = doc.add_table(rows=1, cols=2)
        cat_table.style = 'Light Grid Accent 1'
        h = cat_table.rows[0].cells
        h[0].text = ''
        h[1].text = ''
        _set_cn_font(h[0].paragraphs[0].add_run('类别'))
        _set_cn_font(h[1].paragraphs[0].add_run('数量'))
        for k, v in by_cat.items():
            row = cat_table.add_row().cells
            row[0].text = ''
            row[1].text = ''
            _set_cn_font(row[0].paragraphs[0].add_run(str(k)))
            _set_cn_font(row[1].paragraphs[0].add_run(str(v)))

    # 3. 隐患明细 + 图片
    doc.add_heading('三、隐患明细与图片快照', level=2)
    for idx, det in enumerate(detections, start=1):
        p = doc.add_paragraph()
        _set_cn_font(p.add_run(
            f'{idx}. {det.lab_name or "现场图片"}({det.created_at.strftime("%Y-%m-%d %H:%M")})'))
        img_path = det.annotated_image.path if det.annotated_image else det.original_image.path
        try:
            doc.add_picture(img_path, width=Cm(15))
        except Exception:
            _set_cn_font(doc.add_paragraph().add_run('[图片加载失败]'))
        _set_cn_font(doc.add_paragraph().add_run(f'整体评估:{det.summary or "-"}'))

        if det.hazards:
            tb = doc.add_table(rows=1, cols=6)
            tb.style = 'Light Grid Accent 1'
            headers = ['#', '隐患', '类别', '风险', '描述', '建议']
            for ci, txt in enumerate(headers):
                tb.rows[0].cells[ci].text = ''
                _set_cn_font(tb.rows[0].cells[ci].paragraphs[0].add_run(txt))
            for j, h in enumerate(det.hazards, start=1):
                row = tb.add_row().cells
                vals = [str(j), h.get('name', ''), h.get('category', ''),
                        SEVERITY_LABEL.get(h.get('severity', 'medium'), '中'),
                        h.get('description', ''), h.get('suggestion', '')]
                for ci, v in enumerate(vals):
                    row[ci].text = ''
                    _set_cn_font(row[ci].paragraphs[0].add_run(v), size=9)

    # 4. 参考依据
    if report.references:
        doc.add_heading('四、参考依据', level=2)
        for ref in report.references:
            p = doc.add_paragraph()
            _set_cn_font(p.add_run(f'• {ref.get("title","")}:{ref.get("snippet","")}'))

    # 5. agent 评价
    doc.add_heading('五、AI 评价与整改建议', level=2)
    _add_md_paragraphs(doc, report.agent_evaluation or '(未生成)')

    if report.extra_notes:
        doc.add_heading('六、备注', level=2)
        _set_cn_font(doc.add_paragraph().add_run(report.extra_notes))

    doc.save(str(output_path))
    return output_path
